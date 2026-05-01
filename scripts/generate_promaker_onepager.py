"""
BuildARPro — One-Pager Generator
Generates: D:/Claude Playground/owner_inbox/BuildARPro_OnePager.docx
A4 single page, dense single-column layout, professional pitch format
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"D:\Claude Playground\owner_inbox\BuildARPro_OnePager.docx"

# ─── Colors ───────────────────────────────────────────────────────────────────
NAVY_HEX  = "0A0E1A"
CYAN_HEX  = "00B4D8"
GRAY_HEX  = "444444"
BLACK_HEX = "111111"
LGRAY_HEX = "666666"

def set_font(run, name="Calibri", size_pt=10, bold=False, italic=False,
             color_hex=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)


def set_para_spacing(para, before=0, after=0, line_rule=None, line_val=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def add_heading(doc, text, level_size=13, color=CYAN_HEX, before=6, after=2,
                all_caps=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    run = p.add_run(text.upper() if all_caps else text)
    set_font(run, name="Calibri", size_pt=level_size, bold=True, color_hex=color)
    return p


def add_body(doc, text, size=9.5, color=GRAY_HEX, before=1, after=1,
             italic=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    run = p.add_run(text)
    set_font(run, name="Calibri", size_pt=size, bold=False, italic=italic,
             color_hex=color)
    return p


def add_bullet(doc, text, size=9.5, color=GRAY_HEX):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=0, after=1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, name="Calibri", size_pt=size, color_hex=color)
    return p


def add_divider(doc):
    """Add a thin horizontal rule via a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=3, after=3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00B4D8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_kv_row(doc, key, value, key_size=9.5, val_size=9.5):
    """Bold key: normal value on same line."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    k_run = p.add_run(f"{key}:  ")
    set_font(k_run, size_pt=key_size, bold=True, color_hex=BLACK_HEX)
    v_run = p.add_run(value)
    set_font(v_run, size_pt=val_size, color_hex=GRAY_HEX)
    return p


def set_page_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    """Set page margins in cm."""
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def set_page_a4(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)


def add_table_row(table, cells_data, header=False):
    """Add a row to a table. cells_data = list of (text, width_cm, bold, color)."""
    row = table.add_row()
    for i, (text, w, bold, color) in enumerate(cells_data):
        cell = row.cells[i]
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        set_para_spacing(p, before=1, after=1)
        run = p.add_run(text)
        set_font(run, size_pt=9, bold=bold, color_hex=color)
        if header:
            # shade header cells
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0A0E1A')
            tcPr.append(shd)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_page_a4(doc)
    set_page_margins(doc, top=1.4, bottom=1.4, left=1.8, right=1.8)

    # Remove default styles paragraph spacing
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    nf = style.paragraph_format
    nf.space_before = Pt(0)
    nf.space_after  = Pt(0)

    # ── HEADER ──────────────────────────────────────────────────────────────
    header_p = doc.add_paragraph()
    set_para_spacing(header_p, before=0, after=2)
    title_run = header_p.add_run("BuildARPro")
    set_font(title_run, name="Calibri", size_pt=28, bold=True, color_hex="FFFFFF")
    # set shading on header paragraph to navy
    pPr = header_p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0A0E1A')
    pPr.append(shd)
    header_p.paragraph_format.left_indent = Cm(0.3)

    sub_p = doc.add_paragraph()
    set_para_spacing(sub_p, before=0, after=4)
    sub_run = sub_p.add_run("Building, Reimagined.  |  Build Like a Pro. Every Time.")
    set_font(sub_run, name="Calibri", size_pt=13, bold=False, italic=True,
             color_hex=CYAN_HEX)
    sub_p.paragraph_format.left_indent = Cm(0.3)

    add_divider(doc)

    # ── PROBLEM ─────────────────────────────────────────────────────────────
    add_heading(doc, "The Problem", level_size=11, all_caps=True)
    add_body(doc,
             "600 million people attempt DIY projects every year. Most fail — not because they lack ability, "
             "but because the instructions fail them. Paper manuals are indecipherable. YouTube videos require "
             "constant pausing with dirty hands. Text tutorials leave out the one thing that matters: "
             "where exactly does this piece go?",
             size=9.5)
    add_body(doc,
             "The result: $80 billion wasted annually on failed builds, abandoned projects, and unnecessary "
             "contractor bills. The modern DIYer is motivated — they just need better tools.",
             size=9.5, before=2)

    add_divider(doc)

    # ── SOLUTION ────────────────────────────────────────────────────────────
    add_heading(doc, "The Solution", level_size=11, all_caps=True)
    add_body(doc,
             "BuildARPro is the first platform that turns any paper manual into a live 3D AR guide — "
             "on any smartphone.",
             size=9.5)

    # Re-do with inline bold
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=2)
    r1 = p.add_run("Point your phone camera at a physical instruction sheet. ")
    set_font(r1, size_pt=9.5, color_hex=GRAY_HEX)
    r2 = p.add_run("OCR + Computer Vision + 3D interpretation reads it instantly. ")
    set_font(r2, size_pt=9.5, bold=True, color_hex=BLACK_HEX)
    r3 = p.add_run("A rotatable 3D model appears overlaid on your real world via AR. "
                   "Real-time arrows, highlights, and step-by-step explanations guide every move. "
                   "Your hands stay on the project — not your screen.")
    set_font(r3, size_pt=9.5, color_hex=GRAY_HEX)

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=2, after=2)
    r4 = p2.add_run('"Your IKEA paper manual becomes a live 3D AR guide."')
    set_font(r4, size_pt=10, bold=True, italic=True, color_hex=CYAN_HEX)

    add_divider(doc)

    # ── HOW IT WORKS ────────────────────────────────────────────────────────
    add_heading(doc, "How It Works", level_size=11, all_caps=True)

    steps = [
        ("01 — Choose Your Project",
         "Browse 500+ verified DIY projects by category and skill level. Community ratings tell you what works."),
        ("02 — Scan the Manual",
         "Point camera at any paper instruction sheet. OCR + CV + 3D interpretation activates in seconds."),
        ("03 — Follow the AR Guide",
         "Rotatable 3D model overlays your real space. Arrows and highlights guide every step."),
        ("04 — Complete and Share",
         "Log your build, earn your badge, inspire the next builder."),
    ]
    for num, desc in steps:
        p = doc.add_paragraph()
        set_para_spacing(p, before=1, after=0)
        p.paragraph_format.left_indent = Cm(0.3)
        r_num = p.add_run(f"{num}  ")
        set_font(r_num, size_pt=9.5, bold=True, color_hex=CYAN_HEX)
        r_desc = p.add_run(desc)
        set_font(r_desc, size_pt=9.5, color_hex=GRAY_HEX)

    add_divider(doc)

    # ── MARKET ──────────────────────────────────────────────────────────────
    add_heading(doc, "Market Opportunity", level_size=11, all_caps=True)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    add_table_row(tbl, [
        ("Segment", 5.5, True, "FFFFFF"),
        ("2025 Size", 3.5, True, "FFFFFF"),
        ("CAGR", 2.5, True, "FFFFFF"),
    ], header=True)
    data_rows = [
        ("Global DIY Home Improvement", "$800B", "8%"),
        ("Mobile AR Market", "$30.6B", "31.3%"),
        ("Creator Economy", "$205B", "23.3%"),
    ]
    for seg, size, cagr in data_rows:
        add_table_row(tbl, [
            (seg, 5.5, False, GRAY_HEX),
            (size, 3.5, True, CYAN_HEX),
            (cagr, 2.5, False, GRAY_HEX),
        ])

    add_body(doc,
             "Target Addressable Market: $4–10B near-term | $20B+ at maturity. "
             "2 billion mobile AR users globally — the behavior is mainstream. "
             "The DIY AR application has not been built, until now.",
             size=9.5, before=3)

    add_divider(doc)

    # ── TECHNOLOGY ──────────────────────────────────────────────────────────
    add_heading(doc, "Technology", level_size=11, all_caps=True)

    tech_items = [
        ("OCR Engine", "Scans any paper manual in real-time"),
        ("Computer Vision + 3D", "Interprets 2D diagrams as 3D spatial models"),
        ("ARKit / ARCore", "Precise spatial overlay on iOS + Android"),
        ("Supabase", "Database, Auth, Storage, Edge Functions, Realtime"),
    ]
    for tech, desc in tech_items:
        p = doc.add_paragraph()
        set_para_spacing(p, before=1, after=0)
        p.paragraph_format.left_indent = Cm(0.3)
        rt = p.add_run(f"{tech}: ")
        set_font(rt, size_pt=9.5, bold=True, color_hex=BLACK_HEX)
        rd = p.add_run(desc)
        set_font(rd, size_pt=9.5, color_hex=GRAY_HEX)

    add_divider(doc)

    # ── BUSINESS MODEL ───────────────────────────────────────────────────────
    add_heading(doc, "Business Model", level_size=11, all_caps=True)

    bm_items = [
        ("Consumer Pro — $7.99/month",
         "Unlimited projects, full AI assistant, voice guidance, offline mode, advanced AR."),
        ("Creator Marketplace",
         "Verified creators publish projects, earn revenue share from Pro subscriber completions. Platform: 30%."),
        ("B2B White-Label — $5K–$50K/month",
         "Furniture brands and hardware retailers license BuildARPro to replace paper assembly manuals."),
    ]
    for label, desc in bm_items:
        p = doc.add_paragraph()
        set_para_spacing(p, before=2, after=0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(f"{label}:  ")
        set_font(rl, size_pt=9.5, bold=True, color_hex=CYAN_HEX)
        rd = p.add_run(desc)
        set_font(rd, size_pt=9.5, color_hex=GRAY_HEX)

    add_divider(doc)

    # ── TEAM ────────────────────────────────────────────────────────────────
    add_heading(doc, "Team", level_size=11, all_caps=True)

    p = doc.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    p.paragraph_format.left_indent = Cm(0.3)
    r_name = p.add_run("Inon Baasov")
    set_font(r_name, size_pt=11, bold=True, color_hex=BLACK_HEX)
    r_title = p.add_run("  —  Product Leader & Founder")
    set_font(r_title, size_pt=10, color_hex=CYAN_HEX)

    team_bullets = [
        "Technion (Israel Institute of Technology) graduate",
        "$2.5M+ product impact across prior roles in product leadership",
        "Solo founder, stealth stage, Israel-first market entry",
        "Deep expertise: product strategy, user experience, go-to-market",
    ]
    for b in team_bullets:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        p.paragraph_format.left_indent = Cm(0.7)
        rp = p.add_run(f"•  {b}")
        set_font(rp, size_pt=9.5, color_hex=GRAY_HEX)

    add_divider(doc)

    # ── THE ASK ─────────────────────────────────────────────────────────────
    add_heading(doc, "The Ask", level_size=11, all_caps=True)

    p_ask = doc.add_paragraph()
    set_para_spacing(p_ask, before=1, after=2)
    r_ask = p_ask.add_run(
        "Raising Pre-Seed / Seed capital. "
        "60% engineering (AR + OCR + CV pipeline)  |  "
        "25% creator acquisition + seed content  |  "
        "15% operations."
    )
    set_font(r_ask, size_pt=9.5, color_hex=GRAY_HEX)

    milestones = [
        ("Month 3", "AR MVP live — phone scans any paper manual, 3D AR guide appears in App Store"),
        ("Month 6", "10,000 users, Pro subscription revenue, 50+ verified creator projects"),
        ("Month 9", "First B2B letter of intent from furniture retailer or hardware brand"),
    ]
    for month, goal in milestones:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        p.paragraph_format.left_indent = Cm(0.3)
        rm = p.add_run(f"{month}: ")
        set_font(rm, size_pt=9.5, bold=True, color_hex=CYAN_HEX)
        rg = p.add_run(goal)
        set_font(rg, size_pt=9.5, color_hex=GRAY_HEX)

    # ── FOOTER ──────────────────────────────────────────────────────────────
    # Add actual footer via section header/footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_para.add_run(
        "Confidential  |  Inon Baasov  |  inonbaasov@gmail.com  |  promakerapp.com"
    )
    set_font(r_f, name="Calibri", size_pt=8, color_hex=LGRAY_HEX)

    # ── SAVE ────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) // 1024
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size:  {size_kb} KB ({os.path.getsize(OUTPUT_PATH):,} bytes)")
    print("Done.")


if __name__ == "__main__":
    main()
